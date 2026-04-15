from django.http import HttpResponse
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib import colors
from datetime import date
from django.shortcuts import render, redirect, get_object_or_404
from django.db import IntegrityError
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
import io
import requests
from .models import (
    Member,
    Meeting,
    Attendance,
    WhatsAppDraft,
    Payment,
    Group,
    Topic,
    LessonReflection,
)
from django.core.files import File
from .forms import ReceiptUploadForm
import json
import os
from django.http import HttpResponse, JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.db.models import Count, Sum
from django.utils import timezone
import calendar
from docx import Document

#


def upload_receipt(request):
    error_message = None  # Set a blank error message initially

    if request.method == "POST":
        form = ReceiptUploadForm(request.POST, request.FILES)
        if form.is_valid():
            try:
                # Try to save the form to the database
                form.save()
                return redirect("upload_success")
            except IntegrityError:
                # If the database rejects it because of a duplicate, do this instead:
                error_message = "Hold on! A receipt has already been uploaded for this member for the selected month."
    else:
        form = ReceiptUploadForm()

    return render(
        request, "tracker/upload.html", {"form": form, "error_message": error_message}
    )


def upload_success(request):
    return render(request, "tracker/success.html")


def export_status_pdf(request):
    group_id = request.session.get("active_group_id")
    members = Member.objects.filter(group_id=group_id)

    response = HttpResponse(content_type="application/pdf")
    response["Content-Disposition"] = 'attachment; filename="Roots_Status_Summary.pdf"'

    doc = SimpleDocTemplate(response, pagesize=letter)
    elements = []

    # Table Data Header
    data = [["Member Name", "Total Paid", "Attendance Status"]]

    for member in members:
        # Simplified status for the PDF
        status = "Active" if member.attendance_status != "At Risk" else "Missing"
        data.append(
            [
                f"{member.first_name} {member.last_name}",
                f"${member.total_approved_payments}",
                status,
            ]
        )

    # Styling the Table
    style = TableStyle(
        [
            ("BACKGROUND", (0, 0), (-1, 0), colors.purple),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, 0), 12),
            ("BOTTOMPADDING", (0, 0), (-1, 0), 12),
            ("GRID", (0, 0), (-1, -1), 1, colors.black),
        ]
    )

    table = Table(data)
    table.setStyle(style)
    elements.append(table)
    doc.build(elements)

    return response


def upload_reflections(request):
    group_id = request.session.get("active_group_id") #...
    active_group = Group.objects.get(id=group_id)
    members = Member.objects.filter(group=active_group, is_student=True)
    meetings = Meeting.objects.filter(group=active_group).order_by("-date")

    if request.method == "POST":
        meeting_id = request.POST.get("meeting_id")
        raw_text = request.POST.get("whatsapp_text")
        meeting = get_object_or_404(Meeting, id=meeting_id)

        current_member = None
        reflections_data = {}

        # Read the WhatsApp chat line by line
        # Read the WhatsApp chat line by line
        for line in raw_text.split("\n"):
            line = line.strip()
            if not line:
                continue

            found_member = None
            line_lower = line.lower()

            # SMART MATCHING: Check for the sister's first name in short lines
            # WhatsApp name headers are usually short (under 80 characters)
            if len(line) < 80:
                for member in members:
                    first_name = member.first_name.lower()

                    # If "moreblessing" is in "moreblessing planted"
                    if first_name in line_lower:
                        found_member = member
                        break

            if found_member:
                # We found a new speaker! Switch the lock to her.
                current_member = found_member

                # Create a blank list for her if she hasn't spoken yet
                if current_member.id not in reflections_data:
                    reflections_data[current_member.id] = []

                # We deliberately SKIP appending 'line' here
                # so that "Moreblessing Planted" doesn't get saved as part of her answer!

            elif current_member:
                # If no new name is found, this sentence belongs to the current speaker.
                # We append it to her list.
                reflections_data[current_member.id].append(line)

        # Save all the sorted text into the Database!
        for member_id, lines in reflections_data.items():
            member = members.get(id=member_id)
            combined_text = "\n".join(lines)

            LessonReflection.objects.update_or_create(
                member=member, meeting=meeting, defaults={"my_answers": combined_text}
            )

        return redirect("dashboard")

    return render(
        request,
        "tracker/upload_reflections.html",
        {"meetings": meetings, "active_group": active_group},
    )


def mark_attendance(request):
    group_id = request.session.get("active_group_id")
    if not group_id:
        return redirect("select_group")

    active_group = Group.objects.get(id=group_id)
    members = Member.objects.filter(group=active_group, is_student=True)
    topics = Topic.objects.filter(group=active_group)  # Get all your custom topics

    if request.method == "POST":
        meeting_date = request.POST.get("meeting_date")
        topic_id = request.POST.get("topic_id")  # Get the selected topic ID
        selected_topic = Topic.objects.get(id=topic_id)

        # Create or find the Meeting for this topic/date
        meeting, created = Meeting.objects.get_or_create(
            group=active_group, topic=selected_topic, date=meeting_date
        )

        for member in members:
            mode = request.POST.get(f"mode_{member.id}")
            comment = request.POST.get(f"comment_{member.id}")
            Attendance.objects.update_or_create(
                meeting=meeting,
                member=member,
                defaults={"mode": mode, "comment": comment},
            )
        return redirect("dashboard")

    return render(
        request,
        "tracker/mark_attendance.html",
        {
            "active_group": active_group,
            "members": members,
            "topics": topics,
            "today": timezone.now().date().isoformat(),
        },
    )


def meeting_detail(request, meeting_id):
    meeting = get_object_or_404(Meeting, id=meeting_id)
    attendance_records = Attendance.objects.filter(meeting=meeting)

    # Calculate counts in Python (where math is easy!)
    present_count = attendance_records.filter(mode__in=["physical", "online"]).count()
    absent_count = attendance_records.filter(mode="absent").count()

    if request.method == "POST":
        for record in attendance_records:
            # IMPORTANT: Check your model! Is it record.mode or record.status?
            # If your model uses 'status', change 'mode' below to 'status'
            record.mode = request.POST.get(f"mode_{record.member.id}")
            record.comment = request.POST.get(f"comment_{record.member.id}")
            record.save()
        return redirect("dashboard")

    return render(
        request,
        "tracker/meeting_detail.html",
        {
            "meeting": meeting,
            "attendance_records": attendance_records,
            "present_count": present_count,
            "absent_count": absent_count,
        },
    )


def edit_member(request, member_id):
    member = get_object_or_404(Member, id=member_id)
    if request.method == "POST":
        member.first_name = request.POST.get("first_name")
        member.last_name = request.POST.get("last_name")
        member.phone_number = request.POST.get("phone_number")
        # Add other fields as needed
        member.save()
        return redirect("dashboard")

    return render(request, "tracker/edit_member.html", {"member": member})


def manage_member(request, member_id=None):
    member = get_object_or_404(Member, id=member_id) if member_id else None

    if request.method == "POST":
        first_name = request.POST.get("first_name")
        last_name = request.POST.get("last_name")
        phone = request.POST.get("phone_number")

        if member:  # Updating existing
            member.first_name, member.last_name, member.phone_number = (
                first_name,
                last_name,
                phone,
            )
            member.save()
        else:  # Creating new
            group_id = request.session.get("active_group_id")
            Member.objects.create(
                first_name=first_name,
                last_name=last_name,
                phone_number=phone,
                group_id=group_id,
            )

        return redirect("dashboard")

    return render(request, "tracker/manage_member.html", {"member": member})


def download_summary_summary(request):
    # 1. Get the Active Group
    group_id = request.session.get("active_group_id")
    active_group = Group.objects.get(id=group_id)

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    elements = []
    styles = getSampleStyleSheet()

    # Title with Group Name
    title = Paragraph(f"Roots Summary Report: {active_group.name}", styles["Title"])
    elements.append(title)
    elements.append(Spacer(1, 12))

    # Updated Table Headers
    data = [["Member Name", "Total Paid", "Balance Due", "Meeting & Session Breakdown"]]

    # Get data for this specific group
    all_meetings = Meeting.objects.filter(group=active_group).order_by("date")
    all_members = Member.objects.filter(group=active_group).order_by("first_name")

    # Calculate how many months have passed since the group started (for the $10/month rule)
    # For now, we can count unique months in the meetings list to see how many months they 'should' have paid for
    active_months_count = all_meetings.dates("date", "month").count()
    expected_total = active_months_count * 10

    for member in all_members:
        # Calculate Balance
        total_paid = member.total_approved_payments
        balance = (
            max(0, expected_total - total_paid)
            if not member.is_exempt_from_paying
            else 0
        )

        # Create Meeting Breakdown with Topics
        breakdown_text = ""
        for meeting in all_meetings:
            record = Attendance.objects.filter(member=member, meeting=meeting).first()

            # Show the Topic Name (e.g. Cause & Effect) and the Date
            topic_name = meeting.topic.title if meeting.topic else "General Meeting"
            meeting_label = f"<b>{topic_name}</b> ({meeting.date.strftime('%b %d')})"

            if record and record.mode in ["physical", "online"]:
                status = "<font color='green'>✔ Attended</font>"
            else:
                status = "<font color='red'>✘ Absent</font>"

            breakdown_text += f"{meeting_label}: {status}<br/>"

        if not all_meetings:
            breakdown_text = "No sessions recorded."

        p_breakdown = Paragraph(breakdown_text, styles["Normal"])

        # Add row to PDF
        data.append(
            [
                f"{member.first_name} {member.last_name}",
                f"${total_paid}",
                f"${balance}",
                p_breakdown,
            ]
        )

    # Styling the PDF Table
    t = Table(data, colWidths=[120, 80, 80, 240])
    t.setStyle(
        TableStyle(
            [
                (
                    "BACKGROUND",
                    (0, 0),
                    (-1, 0),
                    colors.HexColor("#8e44ad"),
                ),  # Roots Purple
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("GRID", (0, 0), (-1, -1), 1, colors.grey),
                (
                    "ROWBACKGROUNDS",
                    (0, 1),
                    (-1, -1),
                    [colors.whitesmoke, colors.HexColor("#f4f7f6")],
                ),
            ]
        )
    )

    elements.append(t)
    doc.build(elements)

    pdf = buffer.getvalue()
    buffer.close()
    response = HttpResponse(content_type="application/pdf")
    response["Content-Disposition"] = (
        f'attachment; filename="{active_group.name}_Summary.pdf"'
    )
    response.write(pdf)
    return response


def export_status_pdf(request):
    group_id = request.session.get("active_group_id")
    members = Member.objects.filter(group_id=group_id)

    response = HttpResponse(content_type="application/pdf")
    response["Content-Disposition"] = 'attachment; filename="Roots_Status_Summary.pdf"'

    doc = SimpleDocTemplate(response, pagesize=landscape(letter))
    elements = []

    # Table Headers
    data = [["Member Name", "Total Paid", "Phone Number"]]

    for member in members:
        data.append(
            [
                f"{member.first_name} {member.last_name}",
                f"${member.total_approved_payments}",
                member.phone_number,
            ]
        )

    table = Table(data, colWidths=[200, 150, 200])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#8e44ad")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("BOTTOMPADDING", (0, 0), (-1, 0), 12),
                ("GRID", (0, 0), (-1, -1), 1, colors.black),
            ]
        )
    )

    elements.append(table)
    doc.build(elements)
    return response


def export_status_word(request):
    group_id = request.session.get("active_group_id")
    members = Member.objects.filter(group_id=group_id)

    document = Document()
    document.add_heading("R.O.O.T.S Member Status Summary", 0)

    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Member Name"
    hdr_cells[1].text = "Total Paid"
    hdr_cells[2].text = "Phone Number"

    for member in members:
        row_cells = table.add_row().cells
        row_cells[0].text = f"{member.first_name} {member.last_name}"
        row_cells[1].text = f"${member.total_approved_payments}"
        row_cells[2].text = str(member.phone_number)

    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = 'attachment; filename="Roots_Status_Summary.docx"'
    document.save(response)

    return response


@csrf_exempt
def whatsapp_webhook(request):

    # 1. THE HANDSHAKE
    if request.method == "GET":
        mode = request.GET.get("hub.mode")
        token = request.GET.get("hub.verify_token")
        challenge = request.GET.get("hub.challenge")

        if mode == "subscribe" and token == "roots_secure_token_2026":
            return HttpResponse(challenge, status=200)
        return HttpResponse("error", status=403)

    # 2. THE INBOX & ROUTING
    elif request.method == "POST":
        try:
            body = json.loads(request.body)

            if "object" in body and body["object"] == "whatsapp_business_account":
                entry = body.get("entry", [{}])[0]
                changes = entry.get("changes", [{}])[0]
                value = changes.get("value", {})

                if "messages" in value:
                    message = value["messages"][0]
                    sender_phone = message["from"]
                    message_type = message.get("type")

                    # Route A: It's a standard text message
                    # Route A: It's a standard text message
                    if message_type == "text":
                        incoming_text = message["text"]["body"]
                        print(
                            f"📩 Received TEXT: '{incoming_text}' from {sender_phone}"
                        )
                        send_whatsapp_reply(sender_phone, incoming_text)

                    # Route B: It's an image! (Likely a receipt)
                    elif message_type == "image":
                        image_id = message["image"]["id"]
                        print(
                            f"📸 Received IMAGE with ID: {image_id} from {sender_phone}"
                        )

                        # 🧠 MEMORY STEP 1: Grab the clipboard! Create a fresh draft for this phone number.
                        draft, created = WhatsAppDraft.objects.get_or_create(
                            phone_number=sender_phone
                        )
                        draft.image_id = image_id
                        draft.month = (
                            ""  # Reset the month in case they are starting over
                        )
                        draft.payment_mode = ""  # Reset the payment mode
                        draft.save()

                        # Ask the user if it's a receipt
                        send_receipt_confirmation_button(sender_phone)
                        download_whatsapp_media(image_id)

                    # Route C: It's an interactive button OR list click!
                    elif message_type == "interactive":

                        # 1. Did they click a simple button? (Yes/No OR Payment Mode)
                        if "button_reply" in message["interactive"]:
                            button_id = message["interactive"]["button_reply"]["id"]
                            print(
                                f"🔘 User clicked button ID: {button_id} from {sender_phone}"
                            )

                            if button_id == "btn_yes_receipt":
                                send_month_selection_list(sender_phone)

                            elif button_id == "btn_no_receipt":
                                send_whatsapp_reply(
                                    sender_phone,
                                    "No problem! I will toss that photo in the trash and pretend I didn't see it. 🗑️",
                                )

                            # 🧠 MEMORY STEP 3: The final click! Catch the Payment Mode & Submit to Database.
                            elif button_id in [
                                "btn_pay_cash",
                                "btn_pay_bank",
                                "btn_pay_mobile",
                            ]:
                                draft = WhatsAppDraft.objects.filter(
                                    phone_number=sender_phone
                                ).first()

                                if draft:
                                    # 1. Translate the Payment Mode
                                    if button_id == "btn_pay_cash":
                                        draft.payment_mode = "Cash"
                                        db_method = "physical"
                                    elif button_id == "btn_pay_bank":
                                        draft.payment_mode = "Bank Transfer"
                                        db_method = "bank"
                                    elif button_id == "btn_pay_mobile":
                                        draft.payment_mode = "Mobile Money"
                                        db_method = "ecocash"

                                    # 2. Translate the Month string to a number (e.g., "March" -> 3)
                                    month_map = {
                                        "January": 1,
                                        "February": 2,
                                        "March": 3,
                                        "April": 4,
                                        "May": 5,
                                        "June": 6,
                                        "July": 7,
                                        "August": 8,
                                        "September": 9,
                                        "October": 10,
                                        "November": 11,
                                        "December": 12,
                                    }
                                    db_month = month_map.get(draft.month, 1)

                                    # 3. Find the official Roots Member (matching the last 9 digits of the phone number)
                                    member = Member.objects.filter(
                                        phone_number__endswith=sender_phone[-9:]
                                    ).first()

                                    if member:
                                        try:
                                            # 4. Create the official Payment record!
                                            new_payment = Payment(
                                                member=member,
                                                amount=10.00,
                                                month=db_month,
                                                payment_method=db_method,
                                                status="pending",
                                            )

                                            # 5. Grab the image we downloaded in Route B and attach it
                                            filename = (
                                                f"whatsapp_receipt_{draft.image_id}.jpg"
                                            )
                                            if os.path.exists(filename):
                                                with open(filename, "rb") as f:
                                                    new_payment.receipt_image.save(
                                                        filename, File(f), save=False
                                                    )

                                            new_payment.save()  # Saves to database AND uploads to Cloudinary!

                                            # 6. Success! Send the Boom message.
                                            send_whatsapp_reply(
                                                sender_phone,
                                                f"Boom! 💥 Your $10 {draft.payment_mode} payment receipt for {draft.month} has been successfully submitted to the Roots Command Center for approval. Thank you!",
                                            )

                                            # 7. Clean up: Delete the local image file and wipe the clipboard
                                            if os.path.exists(filename):
                                                os.remove(filename)
                                            draft.delete()

                                        except IntegrityError:
                                            # If they already paid for this month, our database will block it!
                                            send_whatsapp_reply(
                                                sender_phone,
                                                f"Hold on! 🛑 You have already submitted a receipt for {draft.month}. An admin needs to review it first!",
                                            )
                                            draft.delete()  # Wipe the clipboard anyway
                                    else:
                                        send_whatsapp_reply(
                                            sender_phone,
                                            "I received your details, but I couldn't find a Roots member linked to this phone number! Please ask an admin to check your profile. 🕵️‍♀️",
                                        )
                                        draft.delete()
                                else:
                                    send_whatsapp_reply(
                                        sender_phone,
                                        "Oops! I lost your draft. Could you send that photo one more time?",
                                    )
                        # 2. Did they select an item from a list? (The Month)
                        elif "list_reply" in message["interactive"]:
                            list_id = message["interactive"]["list_reply"]["id"]
                            list_title = message["interactive"]["list_reply"][
                                "title"
                            ]  # e.g., "March"
                            print(
                                f"📋 User selected list item: {list_title} ({list_id}) from {sender_phone}"
                            )

                            # 🧠 MEMORY STEP 2: Write down the month they selected!
                            draft = WhatsAppDraft.objects.filter(
                                phone_number=sender_phone
                            ).first()
                            if draft:
                                draft.month = list_title
                                draft.save()

                                # Now ask how they paid!
                                send_payment_mode_buttons(sender_phone)
                return HttpResponse("EVENT_RECEIVED", status=200)

        except Exception as e:
            print(f"❌ Error processing message: {e}")
            return HttpResponse(status=400)

    return HttpResponse("Method not allowed", status=405)


def select_group(request):
    # 1. If they just clicked a group button, save the "sticky note" and let them in!
    if request.method == "POST":
        group_id = request.POST.get("group_id")
        request.session["active_group_id"] = group_id
        return redirect("dashboard")  # We will redirect this to a real Dashboard soon!

    # 2. If they are just arriving, show them all the available groups
    groups = Group.objects.all()
    return render(request, "tracker/select_group.html", {"groups": groups})


def dashboard(request):
    group_id = request.session.get("active_group_id")
    if not group_id:
        return redirect("select_group")

    active_group = Group.objects.get(id=group_id)
    members = Member.objects.filter(group_id=group_id)
    meetings = Meeting.objects.filter(group_id=group_id).order_by("-date")

    # 1. Total Stats
    total_present = Attendance.objects.filter(
        meeting__group_id=group_id, mode__in=["physical", "online"]
    ).count()
    total_absent = Attendance.objects.filter(
        meeting__group_id=group_id, mode="absent"
    ).count()

    # 2. Chart Data
    paid_count = Payment.objects.filter(
        member__group_id=group_id,
        year=date.today().year,
        month=date.today().month,
        status="approved",
    ).count()
    unpaid_count = members.count() - paid_count
    last_5 = meetings[:5]
    meeting_labels = [m.date.strftime("%b %d") for m in last_5]
    attendance_counts = [
        Attendance.objects.filter(meeting=m, mode__in=["physical", "online"]).count()
        for m in last_5
    ]

    # 3. Member payment/attendance breakdown
    current_year = date.today().year
    START_MONTH = 2
    months_to_check = range(START_MONTH, date.today().month + 1)
    latest_meetings = meetings[:10]

    for member in members:
        paid_months_numbers = list(
            Payment.objects.filter(
                member=member, year=current_year, status="approved"
            ).values_list("month", flat=True)
        )
        member.paid_list = [
            calendar.month_abbr[int(m)] for m in paid_months_numbers if m
        ]
        member.unpaid_list = [
            calendar.month_abbr[int(m)]
            for m in months_to_check
            if m not in paid_months_numbers
        ]

        attendance_dict = {
            a.meeting_id: a
            for a in Attendance.objects.filter(
                member=member, meeting__in=latest_meetings
            )
        }
        breakdown = []
        for meeting in latest_meetings:
            record = attendance_dict.get(meeting.id)
            if record:
                status_text = record.get_mode_display()
                color = "red" if record.mode == "absent" else "green"
            else:
                status_text = "Not Marked"
                color = "orange"
            breakdown.append(
                {
                    "date": meeting.date,
                    "topic": meeting.topic.title if meeting.topic else "No Topic",
                    "status": status_text,
                    "color": color,
                    "meeting_id": meeting.id,
                }
            )
        member.recent_breakdown = breakdown

    return render(
        request,
        "tracker/dashboard.html",
        {
            "active_group": active_group,
            "members": members,
            "meetings": meetings,
            "total_present": total_present,
            "total_absent": total_absent,
            "paid_count": paid_count,
            "unpaid_count": unpaid_count,
            "meeting_labels": json.dumps(meeting_labels),
            "attendance_counts": json.dumps(attendance_counts),
        },
    )


# 3. THE OUTBOX (Sending a custom text message back)
def send_whatsapp_reply(recipient_phone, received_text):
    ACCESS_TOKEN = os.getenv("WHATSAPP_ACCESS_TOKEN")
    PHONE_NUMBER_ID = os.getenv("WHATSAPP_PHONE_NUMBER_ID")

    # DEBUG 1: Let's make sure your server can actually see the keys!
    print(f"🔑 DEBUG KEY CHECK - Phone ID: {PHONE_NUMBER_ID}")

    url = f"https://graph.facebook.com/v22.0/{PHONE_NUMBER_ID}/messages"

    headers = {
        "Authorization": f"Bearer {ACCESS_TOKEN}",
        "Content-Type": "application/json",
    }

    reply_text = f"Roots has received your message: '{received_text}'. Your bot is officially live!"

    payload = {
        "messaging_product": "whatsapp",
        "to": recipient_phone,
        "type": "text",
        "text": {"body": reply_text},
    }

    # DEBUG 2: Capture Meta's exact response and print it to the logs
    print(f"🚀 Firing message to {recipient_phone}...")
    response = requests.post(url, headers=headers, json=payload)

    print(f"📤 META API RESPONSE CODE: {response.status_code}")
    print(f"📝 META API DETAILS: {response.text}")


def download_whatsapp_media(media_id):
    ACCESS_TOKEN = os.getenv("WHATSAPP_ACCESS_TOKEN")
    headers = {"Authorization": f"Bearer {ACCESS_TOKEN}"}

    # DANCE STEP 1: Ask Meta for the secure download URL
    print(f"🔍 Asking Meta for the URL for media ID: {media_id}")
    url_request = requests.get(
        f"https://graph.facebook.com/v22.0/{media_id}", headers=headers
    )
    media_data = url_request.json()

    if "url" not in media_data:
        print(f"❌ Failed to get URL. Meta said: {media_data}")
        return None

    media_url = media_data["url"]
    print(f"✅ Got secure URL! Downloading image now...")

    # DANCE STEP 2: Use the URL to download the actual image bytes
    image_response = requests.get(media_url, headers=headers)

    if image_response.status_code == 200:
        # Save it temporarily to your server so we can prove it works!
        filename = f"whatsapp_receipt_{media_id}.jpg"
        with open(filename, "wb") as f:
            f.write(image_response.content)

        print(f"🎉 SUCCESS! Image downloaded and saved as {filename}")
        return filename
    else:
        print(f"❌ Failed to download image. Status: {image_response.status_code}")
        return None


# 4. THE INTERACTIVE MENU (Asking Yes/No)
def send_receipt_confirmation_button(recipient_phone):
    ACCESS_TOKEN = os.getenv("WHATSAPP_ACCESS_TOKEN")
    PHONE_NUMBER_ID = os.getenv("WHATSAPP_PHONE_NUMBER_ID")

    url = f"https://graph.facebook.com/v22.0/{PHONE_NUMBER_ID}/messages"

    headers = {
        "Authorization": f"Bearer {ACCESS_TOKEN}",
        "Content-Type": "application/json",
    }

    # The special payload for Interactive Buttons
    payload = {
        "messaging_product": "whatsapp",
        "to": recipient_phone,
        "type": "interactive",
        "interactive": {
            "type": "button",
            "body": {
                "text": "I just received an image! 📸\n\nIs this a Roots monthly payment receipt?"
            },
            "action": {
                "buttons": [
                    {
                        "type": "reply",
                        "reply": {"id": "btn_yes_receipt", "title": "✅ Yes, it is"},
                    },
                    {
                        "type": "reply",
                        "reply": {
                            "id": "btn_no_receipt",
                            "title": "❌ No, just a photo",
                        },
                    },
                ]
            },
        },
    }

    print(f"🔘 Sending interactive buttons to {recipient_phone}...")
    requests.post(url, headers=headers, json=payload)


# 5. THE LIST MENU (Asking for the Month)
# 5. THE LIST MENU (Asking for the Month)
def send_month_selection_list(recipient_phone):
    ACCESS_TOKEN = os.getenv("WHATSAPP_ACCESS_TOKEN")
    PHONE_NUMBER_ID = os.getenv("WHATSAPP_PHONE_NUMBER_ID")

    url = f"https://graph.facebook.com/v22.0/{PHONE_NUMBER_ID}/messages"

    headers = {
        "Authorization": f"Bearer {ACCESS_TOKEN}",
        "Content-Type": "application/json",
    }

    # WHATSAPP LIMIT: Only 10 rows allowed per list! (We'll use Jan - Oct for now)
    month_rows = [
        {"id": "month_jan", "title": "January"},
        {"id": "month_feb", "title": "February"},
        {"id": "month_mar", "title": "March"},
        {"id": "month_apr", "title": "April"},
        {"id": "month_may", "title": "May"},
        {"id": "month_jun", "title": "June"},
        {"id": "month_jul", "title": "July"},
        {"id": "month_aug", "title": "August"},
        {"id": "month_sep", "title": "September"},
        {"id": "month_oct", "title": "October"},
    ]

    payload = {
        "messaging_product": "whatsapp",
        "to": recipient_phone,
        "type": "interactive",
        "interactive": {
            "type": "list",
            "header": {"type": "text", "text": "📅 Payment Month"},
            "body": {"text": "Perfect! Which month is this $10 receipt paying for?"},
            "footer": {"text": "Roots Command Center"},
            "action": {
                "button": "Tap to Select",
                "sections": [{"title": "2026 Months", "rows": month_rows}],
            },
        },
    }

    print(f"📋 Sending month list menu to {recipient_phone}...")
    response = requests.post(url, headers=headers, json=payload)

    # DEBUG: Let's catch any Meta errors so we aren't blind!
    print(f"📤 LIST MENU RESPONSE CODE: {response.status_code}")
    print(f"📝 LIST MENU DETAILS: {response.text}")


# 6. THE PAYMENT MODE MENU (Cash, Bank, Mobile)
def send_payment_mode_buttons(recipient_phone):
    ACCESS_TOKEN = os.getenv("WHATSAPP_ACCESS_TOKEN")
    PHONE_NUMBER_ID = os.getenv("WHATSAPP_PHONE_NUMBER_ID")

    url = f"https://graph.facebook.com/v22.0/{PHONE_NUMBER_ID}/messages"

    headers = {
        "Authorization": f"Bearer {ACCESS_TOKEN}",
        "Content-Type": "application/json",
    }

    payload = {
        "messaging_product": "whatsapp",
        "to": recipient_phone,
        "type": "interactive",
        "interactive": {
            "type": "button",
            "body": {"text": "Great! 💰 Finally, how was this $10 paid?"},
            "action": {
                "buttons": [
                    {
                        "type": "reply",
                        "reply": {"id": "btn_pay_cash", "title": "💵 Cash"},
                    },
                    {
                        "type": "reply",
                        "reply": {"id": "btn_pay_bank", "title": "🏦 Bank Transfer"},
                    },
                    {
                        "type": "reply",
                        "reply": {"id": "btn_pay_mobile", "title": "📱 Mobile Money"},
                    },
                ]
            },
        },
    }

    print(f"🔘 Sending payment mode buttons to {recipient_phone}...")
    requests.post(url, headers=headers, json=payload)


def member_detail(request, member_id):
    member = get_object_or_404(Member, id=member_id)
    attendances = Attendance.objects.filter(member=member).order_by("-meeting__date")
    payments = Payment.objects.filter(member=member).order_by("-year", "-month")

    # ADD THIS LINE to fetch the sister's journal entries:
    reflections = (
        LessonReflection.objects.filter(member=member)
        .select_related("meeting__topic")
        .order_by("-meeting__date")
    )

    return render(
        request,
        "tracker/member_detail.html",
        {
            "member": member,
            "attendances": attendances,
            "payments": payments,
            "reflections": reflections,  # 👈 Make sure to pass it here!
        },
    )


def write_reflection(request, member_id):
    member = get_object_or_404(Member, id=member_id)
    meetings = Meeting.objects.filter(group=member.group).order_by("-date")

    if request.method == "POST":
        meeting_id = request.POST.get("meeting_id")
        meeting = get_object_or_404(Meeting, id=meeting_id)

        LessonReflection.objects.update_or_create(
            member=member,
            meeting=meeting,
            defaults={
                "my_answers": request.POST.get("answers"),
                "my_commitments": request.POST.get("commitments"),
            },
        )
        return redirect("member_detail", member_id=member.id)

    return render(
        request,
        "tracker/write_reflection.html",
        {"member": member, "meetings": meetings},
    )


def edit_attendance(request, meeting_id):
    # 1. This is the line that went missing! It defines "meeting"
    meeting = get_object_or_404(Meeting, id=meeting_id)

    # 2. Grab all members for this group
    # Change to:
    members = Member.objects.filter(group=meeting.group, is_student=True)

    if request.method == "POST":
        for member in members:
            mode = request.POST.get(f"mode_{member.id}")
            comment_text = request.POST.get(f"comment_{member.id}")

            # Saving with the correct plural 'comments'
            Attendance.objects.update_or_create(
                meeting=meeting,
                member=member,
                defaults={"mode": mode, "comments": comment_text},
            )
        return redirect("dashboard")

    # 3. Pre-load existing records (This is where the crash happened earlier!)
    existing_records = {
        a.member_id: a for a in Attendance.objects.filter(meeting=meeting)
    }

    for member in members:
        record = existing_records.get(member.id)
        member.current_mode = record.mode if record else "absent"
        member.current_comment = record.comments if record else ""

    return render(
        request,
        "tracker/edit_attendance.html",
        {
            "meeting": meeting,
            "members": members,
        },
    )
